import re
from typing import List
from io import BytesIO

import fitz
from docx import Document
from fastapi import FastAPI, UploadFile, File, Form
from pydantic import BaseModel

app = FastAPI(title='AI Resume Analyzer API')

class AnalysisRequest(BaseModel):
    resumeText: str
    jobDescription: str


def _normalize(text: str) -> str:
    return re.sub(r'\s+', ' ', (text or '').strip()).lower()


def _format_label(term: str) -> str:
    mapping = {
        'python': 'Python',
        'sql': 'SQL',
        'machine learning': 'Machine Learning',
        'pandas': 'Pandas',
        'git': 'Git',
        'aws': 'AWS',
        'docker': 'Docker',
        'kubernetes': 'Kubernetes',
        'fastapi': 'FastAPI',
        'linux': 'Linux',
        'classification': 'Classification',
        'tensorflow': 'TensorFlow',
        'deployment': 'Deployment',
        'cloud': 'Cloud',
        'optimization': 'Optimization'
    }
    return mapping.get(term, term.title())


def _extract_skills(text: str) -> List[str]:
    skill_terms = ['python', 'sql', 'machine learning', 'pandas', 'git', 'aws', 'docker', 'kubernetes', 'fastapi', 'linux']
    normalized = _normalize(text)
    found = [_format_label(term) for term in skill_terms if term in normalized]
    return found


def _extract_keywords(text: str) -> List[str]:
    keyword_terms = ['machine learning', 'python', 'classification', 'tensorflow', 'deployment', 'cloud', 'optimization', 'docker']
    normalized = _normalize(text)
    return [_format_label(term) for term in keyword_terms if term in normalized]


def build_analysis_result(resume_text: str, job_description: str):
    resume_norm = _normalize(resume_text)
    jd_norm = _normalize(job_description)

    matched_skills = _extract_skills(resume_text)
    missing_skills = [_format_label(skill) for skill in ['docker', 'aws', 'kubernetes', 'fastapi', 'linux'] if skill.lower() not in resume_norm]
    keywords_present = _extract_keywords(resume_text)
    keywords_missing = [_format_label(keyword) for keyword in ['deployment', 'cloud', 'optimization', 'docker'] if keyword.lower() not in resume_norm]

    ats_score = 70
    if 'python' in resume_norm and 'python' in jd_norm:
        ats_score += 8
    if 'sql' in resume_norm and 'sql' in jd_norm:
        ats_score += 4
    if 'machine learning' in resume_norm and 'machine learning' in jd_norm:
        ats_score += 6
    if 'docker' in resume_norm:
        ats_score += 4
    if 'aws' in resume_norm:
        ats_score += 4
    ats_score = max(60, min(100, ats_score))

    return {
        'atsScore': ats_score,
        'keywordMatch': 88,
        'skillMatch': 82,
        'experienceMatch': 76,
        'educationMatch': 95,
        'formattingScore': 92,
        'readability': 90,
        'matchedSkills': matched_skills,
        'missingSkills': missing_skills,
        'keywordsPresent': keywords_present,
        'keywordsMissing': keywords_missing,
        'strengths': [
            'Strong technical skill section',
            'Good project diversity',
            'Relevant education',
            'Contains measurable achievements'
        ],
        'weaknesses': [
            'Projects lack measurable impact',
            'No certifications included',
            'Experience descriptions are generic',
            'LinkedIn missing'
        ],
        'suggestions': [
            'Add quantified impact in each project bullet.',
            'Tailor your summary to the target role and include the top 3 keywords.',
            'Add a compact certifications section with relevant cloud or platform credentials.'
        ],
        'roadmap': [
            {'priority': 'High', 'items': ['Learn Docker', 'Add FastAPI project', 'Include measurable achievements']},
            {'priority': 'Medium', 'items': ['Add certifications', 'Improve GitHub README']},
            {'priority': 'Low', 'items': ['Portfolio website', 'LinkedIn optimization']}
        ],
        'sectionAnalysis': [
            {'name': 'Contact Information', 'present': 'email' in resume_norm or 'phone' in resume_norm or 'linkedin' in resume_norm},
            {'name': 'Summary', 'present': 'summary' in resume_norm or 'professional summary' in resume_norm},
            {'name': 'Experience', 'present': 'experience' in resume_norm or 'work' in resume_norm},
            {'name': 'Projects', 'present': 'project' in resume_norm},
            {'name': 'Education', 'present': 'education' in resume_norm or 'university' in resume_norm},
            {'name': 'Skills', 'present': 'skills' in resume_norm or 'technical skills' in resume_norm},
            {'name': 'Certifications', 'present': 'certification' in resume_norm},
            {'name': 'Achievements', 'present': 'achievement' in resume_norm or 'impact' in resume_norm}
        ]
    }


@app.get('/health')
def health():
    return {'status': 'ok'}


@app.post('/analyze')
def analyze(req: AnalysisRequest):
    return build_analysis_result(req.resumeText, req.jobDescription)


@app.post('/upload-analyze')
async def upload_analyze(
    resume: UploadFile = File(None),
    job_description: UploadFile = File(None),
    resume_text: str = Form(''),
    jd_text: str = Form('')
):
    text_resume = resume_text
    text_jd = jd_text

    if resume is not None and resume.filename:
        content = await resume.read()
        if resume.filename.lower().endswith('.pdf'):
            doc = fitz.open(stream=content, filetype='pdf')
            text_resume = '\n'.join(page.get_text() for page in doc)
        elif resume.filename.lower().endswith('.docx'):
            document = Document(BytesIO(content))
            text_resume = '\n'.join(p.text for p in document.paragraphs if p.text)

    if job_description is not None and job_description.filename:
        content = await job_description.read()
        if job_description.filename.lower().endswith('.pdf'):
            doc = fitz.open(stream=content, filetype='pdf')
            text_jd = '\n'.join(page.get_text() for page in doc)
        elif job_description.filename.lower().endswith('.docx'):
            document = Document(BytesIO(content))
            text_jd = '\n'.join(p.text for p in document.paragraphs if p.text)

    return build_analysis_result(text_resume, text_jd)
